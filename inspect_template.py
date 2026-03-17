import os
from pathlib import Path
from dotenv import load_dotenv

import sys
sys.path.insert(0, str(Path("execution").resolve()))

from drive_manager import telecharger_modele, MODELE_DOC_ID
from docx import Document

def inspect_placeholders():
    load_dotenv()
    print("Downloading template...")
    chemin_temp = "modele_inspect.docx"
    telecharger_modele(MODELE_DOC_ID, chemin_temp)
    
    print(f"Template downloaded to: {chemin_temp}")
    doc = Document(chemin_temp)
    
    print("--- PARAGRAPHS ---")
    for i, p in enumerate(doc.paragraphs):
        if "{{" in p.text or "}}" in p.text or len(p.text.strip()) > 0:
            print(f"P[{i}]: {p.text}")
            
    print("\n--- TABLES ---")
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, p in enumerate(cell.paragraphs):
                    if "{{" in p.text or "}}" in p.text or len(p.text.strip()) > 0:
                        print(f"T[{t_idx}] R[{r_idx}] C[{c_idx}] P[{p_idx}]: {p.text}")
                        
    if Path(chemin_temp).exists():
        Path(chemin_temp).unlink()

if __name__ == "__main__":
    inspect_placeholders()
