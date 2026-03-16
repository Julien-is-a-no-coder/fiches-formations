"""
Module : google_docs_builder.py (Génération dynamique .docx)
Description : Remplissage d'un fichier .docx local en utilisant python-docx.
Construit la fiche de révision programmatiquement à la suite du modèle.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
import docx.opc.constants
import os
from datetime import datetime
import re

def add_hyperlink(paragraph, url, text, color="2980B9", underline=True):
    """
    Ajoute un lien hypertexte à un paragraphe.
    """
    # Récupérer l'ID de la relation
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Créer l'élément w:hyperlink
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Créer un run pour le texte du lien
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Couleur
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    # Souligné
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    text_element = OxmlElement('w:t')
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def formater_date(date_str: str) -> str:
    """Formate une date en français lisible."""
    mois = {
        "01": "janvier", "02": "février", "03": "mars", "04": "avril",
        "05": "mai", "06": "juin", "07": "juillet", "08": "août",
        "09": "septembre", "10": "octobre", "11": "novembre", "12": "décembre",
    }
    if "-" in date_str and len(date_str) == 10:
        parties = date_str.split("-")
        annee, mois_num, jour = parties[0], parties[1], parties[2]
    elif "/" in date_str:
        parties = date_str.split("/")
        jour, mois_num, annee = parties[0], parties[1], parties[2]
    else:
        return date_str
    return f"{int(jour)} {mois.get(mois_num, mois_num)} {annee}"

def construire_nom_fichier(cursus: str, intitule: str, date: str) -> str:
    if "-" in date and len(date) == 10:
        p = date.split("-")
        date_nom = f"{p[2]}-{p[1]}-{p[0]}"
    else:
        date_nom = date.replace("/", "-")
    intitule_net = re.sub(r'[\\/*?:"<>|]', "-", intitule).strip()
    return f"{cursus}-{intitule_net}-{date_nom}"

def set_cell_background(cell, fill_color):
    """Définit la couleur de fond d'une cellule de tableau (format HEX sans #)."""
    cell_properties = cell._element.get_or_add_tcPr()
    cell_shading = OxmlElement('w:shd')
    cell_shading.set(qn('w:fill'), fill_color)
    cell_properties.append(cell_shading)

def set_table_borders(table):
    """Applique des bordures fines et visibles au tableau."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single') # Ligne simple
        border.set(qn('w:sz'), '4')       # 0.5 point (unité = 1/8 pt)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'BDC3C7') # Gris élégant
        tblBorders.append(border)
    tblPr.append(tblBorders)

def ajouter_titre_section(doc, texte, niveau=1):
    """Ajoute un titre stylisé avec espacement réduit (1 seul saut de ligne)."""
    if niveau == 1:
        # Un seul saut de ligne avant
        doc.add_paragraph()
        
        # Séparateur visuel plus discret
        p_sep = doc.add_paragraph()
        r_sep = p_sep.add_run("________________________________________________________________________________")
        r_sep.font.color.rgb = RGBColor(0xBD, 0xC3, 0xC7)
        r_sep.font.size = Pt(6)
        p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sep.space_after = Pt(12)
        p_sep.space_before = Pt(12)
        
        # Saut de ligne entre le séparateur et le titre
        doc.add_paragraph()
        
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True # Toujours lier au paragraphe suivant
    r = p.add_run(texte)
    r.bold = True
    if niveau == 1:
        r.font.size = Pt(16) # Réduit de 22 à 16
        r.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.space_after = Pt(24)
    elif niveau == 2:
        r.font.size = Pt(14) # Réduit de 16 à 14
        r.font.color.rgb = RGBColor(0x2E, 0x40, 0x53)
        p.space_after = Pt(18)
        p.space_before = Pt(36)
    return p

def ajouter_texte_formate(p, texte, taille=11, couleur=None, gras_global=False):
    """
    Parse le texte pour appliquer le gras (**) et l'italique (*).
    """
    if not texte:
        return
    
    # Regex pour capturer le gras (**), l'italique (*) et les liens [texte](url)
    # On gère d'abord les liens pour éviter les conflits
    link_pattern = r'(\[.*?\]\(.*?\))'
    parts = re.split(link_pattern, texte)
    
    for segment in parts:
        if segment.startswith('[') and '](' in segment and segment.endswith(')'):
            # C'est un lien
            match = re.match(r'\[(.*?)\]\((.*?)\)', segment)
            if match:
                label, url = match.groups()
                # Sécurité : Ne pas ajouter de lien si l'URL est vide ou un simple placeholder
                url_clean = url.strip()
                if url_clean and url_clean not in ["#", "http://", "https://"]:
                    add_hyperlink(p, url_clean, label)
                else:
                    # On affiche le texte seul si le lien est invalide
                    run = p.add_run(label)
                    run.font.size = Pt(taille)
                    if couleur:
                        run.font.color.rgb = couleur
                    run.bold = gras_global
            continue
            
        # Sinon, on parse le gras/italique sur ce segment
        sub_parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', segment)
        for part in sub_parts:
            run = p.add_run()
            run.font.size = Pt(taille)
            if couleur:
                run.font.color.rgb = couleur
            
            if part.startswith('***') and part.endswith('***'):
                run.text = part[3:-3]
                run.bold = True
                run.italic = True
            elif part.startswith('**') and part.endswith('**'):
                run.text = part[2:-2]
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run.text = part[1:-1]
                run.italic = True
            else:
                run.text = part
                run.bold = gras_global

def ajouter_paragraphe(doc, texte, style="Normal", gras=False, puce=None):
    if not texte or texte == "—":
        return
    
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = 1.15
    
    if puce:
        r_puce = p.add_run("• ")
        r_puce.font.size = Pt(11)
        r_puce.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)
        ajouter_texte_formate(p, texte, gras_global=gras)
        p.space_after = Pt(6)
    else:
        ajouter_texte_formate(p, texte, gras_global=gras)
        p.space_after = Pt(10)
    return p

def tenter_remplissage_entete(doc, date_formatee: str, intitule: str, cursus: str):
    """Parcourt tout le document ET ses en-têtes pour remplir les champs dynamiques."""
    remplacements = {
        "{{DATE}}": date_formatee,
        "{{INTITULÉ}}": intitule,
        "{{TITRE}}": intitule,
        "{{CURSUS}}": cursus,
    }

    def traiter_p(p, is_header=False):
        # Remplacement direct placeholders
        for key, val in remplacements.items():
            if key in p.text:
                p.text = p.text.replace(key, str(val))
        
        # Remplacement heuristique "Séance :" ou "Date :"
        txt = p.text.strip()
        if "Séance" in txt and ":" in txt:
            parts = txt.split(":")
            if len(parts) > 1 and len(parts[1].strip()) <= 1:
                p.text = f"Séance : {intitule}"
        if "Date" in txt and ":" in txt:
            parts = txt.split(":")
            if len(parts) > 1 and len(parts[1].strip()) <= 1:
                p.text = f"Date : {date_formatee}"

    # 1. Dans le corps du document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    traiter_p(p)

    # 2. Dans les EN-TÊTES de chaque section
    for section in doc.sections:
        header = section.header
        # Paragraphes simples
        for p in header.paragraphs:
            traiter_p(p, is_header=True)
        # Tableaux d'en-tête (cas des modèles complexes)
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        traiter_p(p, is_header=True)

def remplir_docx_local(
    chemin_modele: str,
    chemin_sortie: str,
    fiche: dict,
    intitule: str,
    date: str,
    cursus: str
) -> bool:
    """
    Construit un fichier .docx de fiche de révision.
    """
    doc = Document(chemin_modele)
    date_formatee = formater_date(date)

    # 1. Remplir l'en-tête du modèle (Toutes pages)
    tenter_remplissage_entete(doc, date_formatee, intitule, cursus)
    
    # 1b. Ajuster les marges pour plus de consistance (Premium look)
    # Réduit de 1.25 à 0.8 pour gagner en espace utile tout en restant élégant
    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    # 1c. Nettoyage du modèle (espace résiduel)
    for p in reversed(doc.paragraphs):
        if not p.text.strip():
            p._element.getparent().remove(p._element)
        else:
            break
            
    # Espace après l'en-tête originel
    p_init = doc.add_paragraph()
    p_init.space_after = Pt(12)
    
    # Espace après l'en-tête originel
    doc.add_paragraph()

    # 2. Titre du document (L'intitulé seul, sur une seule ligne)
    titre_principal = doc.add_paragraph()
    titre_principal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titre = titre_principal.add_run(intitule)
    run_titre.bold = True
    run_titre.font.size = Pt(20) # Réduit de 24 à 20
    run_titre.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
    
    sous_titre = doc.add_paragraph()
    sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sous = sous_titre.add_run(f"{cursus} — Session du {date_formatee}")
    run_sous.italic = True
    run_sous.font.size = Pt(12)
    sous_titre.space_after = Pt(24)

    # 3. L'Essentiel à retenir (Dans un CADRAN / Encadré Premium)
    if fiche.get("l_essentiel"):
        doc.add_paragraph()
        table_ess = doc.add_table(rows=1, cols=1)
        table_ess.width = Inches(6.0)
        table_ess.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Supprimer les bordures par défaut pour un look "flat"
        tbl = table_ess._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
            
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            tblBorders.append(border)
        tblPr.append(tblBorders)

        cell = table_ess.rows[0].cells[0]
        set_cell_background(cell, "F4F6F7") # Gris-bleu très doux
        
        # Contenu de l'encadré
        p_titre = cell.paragraphs[0]
        p_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_titre = p_titre.add_run("L'ESSENTIEL À RETENIR")
        r_titre.bold = True
        r_titre.font.size = Pt(10)
        r_titre.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
        p_titre.space_before = Pt(12)
        p_titre.space_after = Pt(6)
        
        p_txt = cell.add_paragraph()
        p_txt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ajouter_texte_formate(p_txt, fiche["l_essentiel"], taille=11, couleur=RGBColor(0x2C, 0x3E, 0x50))
        p_txt.space_after = Pt(12)
        p_txt.paragraph_format.left_indent = Pt(20)
        p_txt.paragraph_format.right_indent = Pt(20)

        doc.add_paragraph().space_after = Pt(24) # Espace après l'encadré

    # 4. Objectifs de la séance
    if fiche.get("les_objectifs"):
        ajouter_titre_section(doc, "Compétences Visées", niveau=2)
        for obj in fiche["les_objectifs"]:
            ajouter_paragraphe(doc, obj, puce=True)

    # 5. Sections Principales
    for i, section_data in enumerate(fiche.get("sections_principales", [])):
        doc.add_paragraph() # Saut de ligne avant chaque nouvelle section principale
        ajouter_titre_section(doc, section_data["titre"], niveau=1)
        
        # Chapeau introductif
        if section_data.get("chapeau_introductif"):
            p_chap = ajouter_paragraphe(doc, section_data["chapeau_introductif"])
            if p_chap:
                p_chap.runs[0].italic = True
                p_chap.paragraph_format.space_after = Pt(12)
            doc.add_paragraph() # Saut de ligne après le chapeau
        
        # Points clés
        for pt in section_data.get("points_cles", []):
            ajouter_paragraphe(doc, pt, puce=True)
        
        doc.add_paragraph() # Saut de ligne avant le tableau ou conseil
        # Conseil de l'expert (S'il existe pour cette section)
        conseil = section_data.get("conseil_expert")
        if conseil:
            doc.add_paragraph() # Saut avant le conseil
            p_conseil = doc.add_paragraph()
            # On sépare l'icône du texte pour formater le reste
            p_conseil.add_run("💡 Conseil : ").italic = True
            ajouter_texte_formate(p_conseil, conseil, taille=10, couleur=RGBColor(0x29, 0x80, 0xB9))
            p_conseil.space_before = Pt(6)
            p_conseil.space_after = Pt(6)
            doc.add_paragraph() # Saut après le conseil

        # Si un tableau est présent et demandé
        tab_data = section_data.get("tableau_comparatif") or {}
        if tab_data.get("afficher") and tab_data.get("entetes") and tab_data.get("lignes"):
            # Titre du tableau
            titre_tab = tab_data.get("titre_tableau", "")
            if titre_tab:
                p_tab = doc.add_paragraph()
                p_tab.paragraph_format.keep_with_next = True
                r_tab = p_tab.add_run(titre_tab)
                r_tab.bold = True
                r_tab.font.size = Pt(12)
                p_tab.space_before = Pt(30)
                p_tab.space_after = Pt(12)
                
            nb_cols = len(tab_data["entetes"])
            nb_lignes = len(tab_data["lignes"])
            table = doc.add_table(rows=nb_lignes + 1, cols=nb_cols)
            # Appliquer les bordures fines personnalisées au lieu du style par défaut
            set_table_borders(table)
            
            # EMPÊCHER LA COUPURE DU TABLEAU SUR DEUX PAGES
            # On demande à Word de ne pas couper les lignes individuelles
            for row in table.rows:
                tr = row._tr
                trPr = tr.get_or_add_trPr()
                cantSplit = OxmlElement('w:cantSplit')
                cantSplit.set(qn('w:val'), 'true')
                trPr.append(cantSplit)
            
            # Espacement dans les cellules
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].paragraph_format.space_before = Pt(6)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(6)

            # Entêtes
            hdr_cells = table.rows[0].cells
            for j, entete in enumerate(tab_data["entetes"]):
                p_cell = hdr_cells[j].paragraphs[0]
                ajouter_texte_formate(p_cell, entete, taille=10, gras_global=True)
                p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Lignes
            for row_idx, ligne_donnees in enumerate(tab_data["lignes"]):
                row_cells = table.rows[row_idx + 1].cells
                for col_idx, valeur in enumerate(ligne_donnees):
                    if col_idx < nb_cols:
                        p_cell = row_cells[col_idx].paragraphs[0]
                        ajouter_texte_formate(p_cell, str(valeur), taille=10)
            
            doc.add_paragraph().space_after = Pt(18) # Espace après le tableau

    # 6. Cas Pratique / Atelier
    cas_prat = fiche.get("cas_pratique") or {}
    if cas_prat.get("afficher") and cas_prat.get("situations"):
        doc.add_paragraph()
        doc.add_paragraph()
        ajouter_titre_section(doc, cas_prat.get("titre_atelier", "ATELIER PRATIQUE"), niveau=1)
        doc.add_paragraph()

        for d_p in cas_prat.get("organisation_livrables", []):
            ajouter_paragraphe(doc, d_p, puce=True)
        
        doc.add_paragraph()
        for sit in cas_prat.get("situations", []):
            p_sit = doc.add_paragraph()
            p_sit.paragraph_format.line_spacing = 1.15
            run_nom = p_sit.add_run(f"{sit.get('nom_situation', 'Cas')} : ")
            run_nom.bold = True
            ajouter_texte_formate(p_sit, sit.get('faits_et_attendus', ''))
            p_sit.space_after = Pt(8)

    # 7. Points de Vigilance et Ressources
    # Saut de ligne demandé avant la vigilance
    doc.add_paragraph()
    doc.add_paragraph()
    
    vigilance_all = fiche.get("points_vigilance_et_ressources") or []
    if vigilance_all:
        vigilance_points = [v for v in vigilance_all if "Vigilance" in v or "Attention" in v]
        ressources_points = [v for v in vigilance_all if "Ressource" in v or "Lien" in v or "Fichier" in v or "PDF" in v]
        
        # S'il n'y a pas de distinction claire, on met tout dans vigilance par défaut
        if not vigilance_points and not ressources_points:
            vigilance_points = vigilance_all

        # --- POINTS DE VIGILANCE ---
        if vigilance_points:
            ajouter_titre_section(doc, "Points de Vigilance", niveau=2)
            for v in vigilance_points:
                # Nettoyer le préfixe si présent pour plus de propreté
                v_clean = v.replace("Vigilance :", "").replace("Vigilance:", "").strip()
                ajouter_paragraphe(doc, v_clean, puce=True)

        # --- CADRAN RESSOURCES ---
        if ressources_points:
            doc.add_paragraph()
            doc.add_paragraph()
            
            table_res = doc.add_table(rows=1, cols=1)
            table_res.width = Inches(6.0)
            table_res.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Supprimer les bordures
            tbl = table_res._tbl
            tblPr = tbl.tblPr
            if tblPr is not None:
                tblBorders = OxmlElement('w:tblBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    tblBorders.append(border)
                tblPr.append(tblBorders)

            cell = table_res.rows[0].cells[0]
            set_cell_background(cell, "EBF5FB") # Bleu très clair pour changer du gris
            
            p_titre = cell.paragraphs[0]
            p_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_titre = p_titre.add_run("POUR ALLER PLUS LOIN : RESSOURCES & OUTILS")
            r_titre.bold = True
            r_titre.font.size = Pt(10)
            r_titre.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
            p_titre.space_before = Pt(12)
            p_titre.space_after = Pt(6)
            
            for res in ressources_points:
                res_clean = res.replace("Ressource :", "").replace("Ressource:", "").strip()
                p_res = cell.add_paragraph()
                p_res.paragraph_format.left_indent = Pt(30)
                r_bullet = p_res.add_run("🔗 ")
                ajouter_texte_formate(p_res, res_clean, taille=10, couleur=RGBColor(0x21, 0x61, 0x8C))
                p_res.space_after = Pt(4)

            # Espace de fin dans le cadran
            doc.add_paragraph().space_after = Pt(6)

    # Fin du document
    doc.add_paragraph()
    doc.add_paragraph()
    doc.save(chemin_sortie)
    return True
